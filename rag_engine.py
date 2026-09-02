"""محرّك RAG — منفصل تماماً عن الواجهة."""

import hashlib
import io
import os

import chromadb
import docx
import fitz
import httpx
import openpyxl
from chromadb.utils import embedding_functions
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter

load_dotenv()

GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
MODEL = os.getenv("LLM_MODEL", "openai/gpt-oss-120b")
EMBED_MODEL = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
VECTORSTORE_PATH = "./vectorstore"
COLLECTION_NAME = "company_docs"

CHUNK_SIZE = 500
CHUNK_OVERLAP = 100
MAX_DISTANCE = 0.55
TOP_K = 3

SYSTEM_PROMPT = """أنت مساعد مؤسسي. أجب على السؤال بالاعتماد على السياق المرفق فقط.

قواعد إلزامية:
- لا تستخدم أي معرفة خارج السياق المرفق.
- إذا لم يحتوِ السياق على الإجابة، قل: "لا تحتوي المستندات المتاحة على إجابة لهذا السؤال."
- لا تخمّن ولا تستنتج معلومات غير مذكورة صراحة.
- أشر إلى رقم المصدر الذي اعتمدت عليه، مثل [مصدر 1].
- أجب بإيجاز."""


# ---------- قاعدة البيانات ----------

def get_collection():
    client = chromadb.PersistentClient(path=VECTORSTORE_PATH)
    embed_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=EMBED_MODEL
    )
    return client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=embed_fn,
        metadata={"hnsw:space": "cosine"},
    )


# ---------- استخراج النص ----------

def extract_pdf(file_bytes: bytes) -> list[tuple[int, str, str]]:
    """يعيد (رقم تسلسلي، وصف الموقع، نص)."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    out = [(i + 1, f"صفحة {i + 1}", page.get_text()) for i, page in enumerate(doc)]
    doc.close()
    return out


def extract_docx(file_bytes: bytes) -> list[tuple[int, str, str]]:
    d = docx.Document(io.BytesIO(file_bytes))

    parts = [p.text for p in d.paragraphs if p.text.strip()]

    for t_idx, table in enumerate(d.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
        ]
        parts.append(f"[جدول {t_idx}]\n" + "\n".join(rows))

    full_text = "\n".join(parts)
    return [(1, "المستند", full_text)] if full_text.strip() else []


def extract_xlsx(file_bytes: bytes) -> list[tuple[int, str, str]]:
    """كل صف قطعة مستقلة. يتوقف عند أول عمود فارغ في صف العناوين."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    out = []

    for idx, sheet in enumerate(wb.worksheets, start=1):
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        # نقتصر على الأعمدة المتصلة من بداية صف العناوين
        headers = []
        for c in rows[0]:
            if c is None or not str(c).strip():
                break
            headers.append(str(c).strip())

        if not headers:
            continue

        n = len(headers)
        for row_num, row in enumerate(rows[1:], start=2):
            pairs = [
                f"{headers[i]}: {row[i]}"
                for i in range(min(n, len(row)))
                if row[i] is not None
            ]
            if pairs:
                text = f"[{sheet.title}] " + "، ".join(pairs)
                out.append((idx, f"ورقة: {sheet.title}، صف {row_num}", text))

    wb.close()
    return out

EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "xlsx": extract_xlsx,
}

SUPPORTED_TYPES = list(EXTRACTORS.keys())


def extract_any(file_bytes: bytes, filename: str) -> list[tuple[int, str, str]]:
    ext = filename.rsplit(".", 1)[-1].lower()
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(f"صيغة غير مدعومة: .{ext}")
    return extractor(file_bytes)


# ---------- التقطيع ----------

def chunk_sections(sections: list[tuple[int, str, str]]) -> list[tuple[int, str, str]]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", "، ", " ", ""],
    )
    out = []
    for num, location, text in sections:
        if not text.strip():
            continue
        for chunk in splitter.split_text(text):
            out.append((num, location, chunk))
    return out


# ---------- إدارة المستندات ----------

def file_hash(file_bytes: bytes) -> str:
    return hashlib.md5(file_bytes).hexdigest()[:12]


def document_exists(collection, doc_hash: str) -> bool:
    res = collection.get(where={"doc_hash": doc_hash}, limit=1)
    return len(res["ids"]) > 0


def ingest_file(collection, file_bytes: bytes, filename: str) -> dict:
    doc_hash = file_hash(file_bytes)

    if document_exists(collection, doc_hash):
        return {"status": "duplicate", "filename": filename, "chunks": 0}

    sections = extract_any(file_bytes, filename)
    chunks = chunk_sections(sections)

    if not chunks:
        return {"status": "empty", "filename": filename, "chunks": 0}

    collection.add(
        documents=[c[2] for c in chunks],
        metadatas=[
            {
                "source": filename,
                "page": c[0],
                "location": c[1],
                "doc_hash": doc_hash,
            }
            for c in chunks
        ],
        ids=[f"{doc_hash}_{i}" for i in range(len(chunks))],
    )
    return {"status": "added", "filename": filename, "chunks": len(chunks)}


def delete_document(collection, doc_hash: str) -> None:
    collection.delete(where={"doc_hash": doc_hash})


def list_documents(collection) -> dict[str, dict]:
    res = collection.get(include=["metadatas"])
    docs = {}
    for meta in res["metadatas"]:
        h = meta.get("doc_hash", "legacy")
        if h not in docs:
            docs[h] = {"filename": meta.get("source", "غير معروف"), "chunks": 0}
        docs[h]["chunks"] += 1
    return docs


# ---------- الاسترجاع والتوليد ----------

def retrieve(collection, question: str) -> list[dict]:
    if collection.count() == 0:
        return []
    res = collection.query(query_texts=[question], n_results=TOP_K)
    chunks = [
        {"text": d, "meta": m, "distance": dist}
        for d, m, dist in zip(
            res["documents"][0], res["metadatas"][0], res["distances"][0]
        )
    ]
    return [c for c in chunks if c["distance"] <= MAX_DISTANCE]


def build_context(chunks: list[dict]) -> str:
    return "\n\n".join(
        f"[مصدر {i}] (ملف: {c['meta']['source']}، {c['meta'].get('location', '')})\n{c['text']}"
        for i, c in enumerate(chunks, start=1)
    )


def ask_rag(collection, question: str) -> dict:
    chunks = retrieve(collection, question)

    if not chunks:
        return {
            "answer": "لا تحتوي المستندات المتاحة على معلومات ذات صلة بهذا السؤال.",
            "sources": [],
        }

    r = httpx.post(
        GROQ_URL,
        headers={"Authorization": f"Bearer {os.getenv('GROQ_API_KEY')}"},
        json={
            "model": MODEL,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": f"السياق:\n{build_context(chunks)}\n\nالسؤال: {question}",
                },
            ],
            "temperature": 0.0,
        },
        timeout=60,
    )
    if r.status_code != 200:
        raise RuntimeError(f"LLM error {r.status_code}: {r.text}")

    return {
        "answer": r.json()["choices"][0]["message"]["content"],
        "sources": chunks,
    }